from flask import Blueprint, render_template
from flask_login import login_required, current_user
from flask import Flask, request, render_template, send_file
import pandas as pd
from docx.shared import Inches
from io import BytesIO
from docx import Document
from docxtpl import DocxTemplate
from docx2pdf import convert
from fpdf import FPDF
import base64


views = Blueprint("views", __name__)


@views.route("/")
@login_required
def home():
    return render_template("home.html", user=current_user)


@views.route("/instructions")
@login_required
def question_paper():
    return render_template("instructions.html", user=current_user)


@views.route("/generate", methods=["GET", "POST"])
@login_required
def generate_question_paper():
    file = request.files['excel_file']
    df = pd.read_excel(file)
    question = df["Question"]
    marks = df["Marks"]
    difficulty = df["Difficulty"]
    blooms = df["Blooms"]
    # Test hard code
    df = df[((df['Marks'] == 5) & (df['Difficulty'] == "Easy")) |
            ((df['Marks'] == 3) & (df['Difficulty'] == "Medium"))]
    df = df.sample(n=5, replace=False)
    uniqueQuestionList = df["Question"].to_list()
    marksList = df["Marks"].to_list()

    # document = Document()

    # p1 = document.add_paragraph()
    # p1.alignment = 1
    # run = p1.add_run("MIT ADT UNIVERSITY PAPER")
    # run.bold = True
    # table = document.add_table(rows=7, cols=3)

    # for i in range(0, 4):
    #     for j in range(0, 3):

    #         questionNoCell = table.cell(i, 0)
    #         questionNoCell.text = (f"Q{i+1})")
    #         questionNoCell.width = Inches(0.2)
    #         questionCell = table.cell(i, 1)
    #         questionCell.text = uniqueQuestionList[i]    # Error
    #         questionCell.width = Inches(60)
    #         marksCell = table.cell(i, 2)
    #         marksCell.text = str(marksList[i])
    #         marksCell.width = Inches(0.2)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=14)
    pdf.cell(200, 10, txt="MIT ADT UNIVERSITY PAPER", align="C")
    pdf.ln()
    pdf.set_font("Arial", size=12)
    for i in range(0, 5):
        pdf.cell(10, 10, txt=f"Q{i+1}", ln=0)
        pdf.multi_cell(
            160, 10, txt=f"{uniqueQuestionList[i]} ({marksList[i]})")

    pdf_bytes = pdf.output(dest='S').encode('latin1')
    pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')

    return render_template("generate.html", user=current_user, document_pdf=pdf_base64, df=df)
