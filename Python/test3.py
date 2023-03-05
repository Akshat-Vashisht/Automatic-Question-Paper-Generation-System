from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx import Document
import pandas as pd
import random
from docx2pdf import convert
from docx.shared import Inches
import os


# def insertHR(paragraph):
#     p = paragraph._p
#     pPr = p.get_or_add_pPr()
#     pBdr = OxmlElement('w:pBdr')
#     pPr.insert_element_before(pBdr,
#                               'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
#                               'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
#                               'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
#                               'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
#                               'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
#                               'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
#                               'w:pPrChange'
#                               )
#     bottom = OxmlElement('w:bottom')
#     bottom.set(qn('w:val'), 'single')
#     bottom.set(qn('w:sz'), '6')
#     bottom.set(qn('w:space'), '1')
#     bottom.set(qn('w:color'), 'auto')
#     pBdr.append(bottom)


df = pd.read_excel('TestQuestionBank.xlsx')

df = df[(df['Marks'] == 5) & (df['Difficulty'] == "Easy")]

listOf5 = df['Question'].to_list()

UniqueListOf5 = list()
i = 0
for i in range(0, 5):
    q = random.choice(listOf5)
    if q in UniqueListOf5:
        i = i - 1
    else:
        UniqueListOf5.append(q)

for i in UniqueListOf5:
    print(i)

document = Document()

p1 = document.add_paragraph()
p1.alignment = 1
run = p1.add_run("MIT ADT UNIVERSITY PAPER")
run.bold = True
# insertHR(p1)

table = document.add_table(rows=7, cols=3)

for i in range(0, 4):
    for j in range(0, 3):

        questionNoCell = table.cell(i, 0)
        questionNoCell.text = (f"Q{i+1})")
        questionNoCell.width = Inches(0.2)
        questionCell = table.cell(i, 1)
        questionCell.text = UniqueListOf5[i]    # Error
        questionCell.width = Inches(60)
        marksCell = table.cell(i, 2)
        marksCell.text = "(5)"
        marksCell.width = Inches(0.2)


document.save("Test.docx")
convert("Test.docx")
os.system("Test.pdf")
